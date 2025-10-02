"""
Word Document (DOCX) Export Utilities
Handles exporting OCR text and structured template data to Microsoft Word format
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os


class DocxExporter:
    """Export OCR data to Word documents with professional formatting"""
    
    def __init__(self):
        self.default_font_size = Pt(11)
        self.heading_font_size = Pt(14)
        self.title_font_size = Pt(16)
    
    def export_plain_text_to_docx(self, text, output_path, title="", metadata=None):
        """
        Export plain text to a Word document
        
        Args:
            text: The text content to export
            output_path: Path where the DOCX file will be saved
            title: Optional document title
            metadata: Optional dict with author, date, confidence, etc.
        
        Returns:
            Path to the created DOCX file
        """
        try:
            # Create new document
            doc = Document()
            
            # Set default styles
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = self.default_font_size
            
            # Add title if provided
            if title:
                title_para = doc.add_heading(title, level=0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata section if provided
            if metadata:
                meta_para = doc.add_paragraph()
                meta_para.add_run("Document Information").bold = True
                meta_para.add_run("\n")
                
                for key, value in metadata.items():
                    meta_para.add_run(f"{key}: ").bold = True
                    meta_para.add_run(f"{value}\n")
                
                # Add separator
                doc.add_paragraph("_" * 80)
                doc.add_paragraph()
            
            # Add main content
            # Split text into paragraphs and add each
            paragraphs = text.split('\n\n') if '\n\n' in text else text.split('\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text.strip())
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add footer with export timestamp
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"Exported on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save document
            doc.save(output_path)
            return output_path
            
        except Exception as e:
            raise Exception(f"Error creating DOCX file: {str(e)}")
    
    def export_template_data_to_docx(self, document_data, template, output_path, title=""):
        """
        Export structured template data to a formatted Word document
        
        Args:
            document_data: Extracted data dictionary from Document model
            template: Template model instance
            output_path: Path where the DOCX file will be saved
            title: Optional document title
        
        Returns:
            Path to the created DOCX file
        """
        try:
            # Create new document
            doc = Document()
            
            # Add title
            if not title and hasattr(template, 'name'):
                title = f"Extracted Data - {template.name}"
            
            if title:
                title_heading = doc.add_heading(title, level=0)
                title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add document metadata
            doc.add_paragraph()
            meta_table = doc.add_table(rows=3, cols=2)
            meta_table.style = 'Light Grid Accent 1'
            
            # Template info
            meta_table.cell(0, 0).text = "Template"
            meta_table.cell(0, 1).text = template.name if hasattr(template, 'name') else "N/A"
            
            # Processing date
            meta_table.cell(1, 0).text = "Processed Date"
            meta_table.cell(1, 1).text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            
            # Field count
            field_count = len(document_data.get('fields', [])) or len(document_data.get('cells', []))
            meta_table.cell(2, 0).text = "Fields Extracted"
            meta_table.cell(2, 1).text = str(field_count)
            
            doc.add_paragraph()
            doc.add_heading("Extracted Fields", level=1)
            
            # Extract fields based on format
            if 'fields' in document_data:
                # Template-based format
                fields = document_data['fields']
                
                # Create table for fields
                if fields:
                    # Add 3 columns: Field Name, Value, Confidence
                    field_table = doc.add_table(rows=1, cols=3)
                    field_table.style = 'Light Grid Accent 1'
                    
                    # Header row
                    header_cells = field_table.rows[0].cells
                    header_cells[0].text = "Field Name"
                    header_cells[1].text = "Value"
                    header_cells[2].text = "Confidence"
                    
                    # Make header bold
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    # Add data rows
                    for field in fields:
                        row_cells = field_table.add_row().cells
                        row_cells[0].text = field.get('name', '')
                        row_cells[1].text = field.get('value', '')
                        confidence = field.get('confidence', 0)
                        row_cells[2].text = f"{confidence:.1f}%" if confidence else "N/A"
                
            elif 'cells' in document_data:
                # Table detection format
                cells = document_data['cells']
                headers = document_data.get('headers', {})
                rows_count = document_data.get('rows', 0)
                cols_count = document_data.get('cols', 0)
                
                if cells and cols_count > 0:
                    # Group cells by row
                    row_data = {}
                    for cell in cells:
                        row = cell.get('row', 0)
                        col = cell.get('col', 0)
                        text = cell.get('text', '')
                        
                        if row not in row_data:
                            row_data[row] = {}
                        row_data[row][col] = text
                    
                    # Create table
                    data_table = doc.add_table(rows=1, cols=cols_count)
                    data_table.style = 'Light Grid Accent 1'
                    
                    # Add headers
                    header_cells = data_table.rows[0].cells
                    for col_idx in range(cols_count):
                        header_text = headers.get(str(col_idx), f"Column {col_idx + 1}")
                        header_cells[col_idx].text = header_text
                        for paragraph in header_cells[col_idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    # Add data rows (skip header row 0)
                    for row_idx in sorted(row_data.keys()):
                        if row_idx == 0:  # Skip header row
                            continue
                        
                        row_cells = data_table.add_row().cells
                        for col_idx in range(cols_count):
                            cell_text = row_data[row_idx].get(col_idx, '')
                            row_cells[col_idx].text = cell_text
            
            # Add footer
            doc.add_paragraph()
            doc.add_paragraph("_" * 80)
            footer_para = doc.add_paragraph()
            footer_para.add_run("Generated by OCR App\n").italic = True
            footer_para.add_run(f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save document
            doc.save(output_path)
            return output_path
            
        except Exception as e:
            raise Exception(f"Error creating template DOCX: {str(e)}")
    
    def export_multi_documents_to_docx(self, documents, template, output_path, title=""):
        """
        Export multiple documents to a single consolidated Word document
        
        Args:
            documents: List of Document model instances
            template: Template model instance
            output_path: Path where the DOCX file will be saved
            title: Optional document title
        
        Returns:
            Path to the created DOCX file
        """
        try:
            # Create new document
            doc = Document()
            
            # Add title
            if not title:
                title = f"Consolidated Data - {template.name}"
            
            title_heading = doc.add_heading(title, level=0)
            title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add summary
            doc.add_paragraph()
            summary_para = doc.add_paragraph()
            summary_para.add_run(f"Template: ").bold = True
            summary_para.add_run(f"{template.name}\n")
            summary_para.add_run(f"Total Documents: ").bold = True
            summary_para.add_run(f"{len(documents)}\n")
            summary_para.add_run(f"Export Date: ").bold = True
            summary_para.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            
            doc.add_paragraph()
            doc.add_heading("Extracted Data", level=1)
            
            # Get field names from template
            field_names = template.get_field_names() if hasattr(template, 'get_field_names') else []
            
            if field_names:
                # Create consolidated table
                table = doc.add_table(rows=1, cols=len(field_names) + 1)
                table.style = 'Light Grid Accent 1'
                
                # Header row
                header_cells = table.rows[0].cells
                header_cells[0].text = "Document"
                for idx, field_name in enumerate(field_names, start=1):
                    header_cells[idx].text = field_name
                
                # Make header bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Add data for each document
                for doc_idx, document in enumerate(documents, start=1):
                    row_cells = table.add_row().cells
                    row_cells[0].text = f"Doc {doc_idx}"
                    
                    extracted_data = document.extracted_data
                    
                    if 'fields' in extracted_data:
                        fields = extracted_data['fields']
                        for idx, field in enumerate(fields):
                            if idx + 1 < len(row_cells):
                                row_cells[idx + 1].text = field.get('value', '')
                    
                    elif 'cells' in extracted_data:
                        cells = extracted_data['cells']
                        # Extract first data row
                        row_data = {}
                        for cell in cells:
                            row = cell.get('row', 0)
                            col = cell.get('col', 0)
                            if row == 1:  # First data row
                                row_data[col] = cell.get('text', '')
                        
                        for col_idx in sorted(row_data.keys()):
                            if col_idx + 1 < len(row_cells):
                                row_cells[col_idx + 1].text = row_data[col_idx]
            
            # Add footer
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"Generated by OCR App | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save document
            doc.save(output_path)
            return output_path
            
        except Exception as e:
            raise Exception(f"Error creating multi-document DOCX: {str(e)}")
